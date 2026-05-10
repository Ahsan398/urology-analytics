import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_documentation():
    doc = Document()
    
    # --- Styling ---
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Urology Analytics System - Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Comprehensive Overview, Architecture, and LLM Integration Details")
    doc.add_paragraph("_" * 70)
    
    # --- Section 1: How the Project Started ---
    doc.add_heading('1. Project Origins and Objectives', level=1)
    p = doc.add_paragraph(
        "This project was initiated to build a robust, end-to-end data analytics pipeline tailored for a Urology Practice, "
        "specifically designed as a portfolio piece for a Johns Hopkins Hospital operations analytics role. "
        "The primary goal was to move beyond simple spreadsheets and demonstrate production-level data engineering, predictive modeling, "
        "and automated reporting. "
    )
    p = doc.add_paragraph(
        "By ingesting massive, real-world public healthcare datasets (CMS Medicare Provider Utilization, HCUP, MEPS, and Hospital Compare), "
        "the project establishes a defensible and highly professional 'single source of truth' for analyzing physician productivity, billing compliance, "
        "and clinical capacity."
    )
    
    # --- Section 2: Files and Scripts Overview ---
    doc.add_heading('2. System Architecture & Scripts Pipeline', level=1)
    doc.add_paragraph("The pipeline consists of 17 sequential Python scripts, structured systematically into modules:")
    
    # Module 1: Setup & Ingestion
    doc.add_heading('Module 1: Setup & Data Ingestion', level=2)
    doc.add_paragraph("• 00_setup_environment.py: Initializes the directory structure (data/, outputs/, logs/) to ensure reproducible execution.", style='List Bullet')
    doc.add_paragraph("• 01_download_cms_data.py: Ingests Medicare Physician and Other Practitioner data.", style='List Bullet')
    doc.add_paragraph("• 02_download_hcup_data.py: Simulates downloading Healthcare Cost and Utilization Project data for procedure trends.", style='List Bullet')
    doc.add_paragraph("• 03_download_meps_data.py: Ingests Medical Expenditure Panel Survey data to gauge patient access and wait times.", style='List Bullet')
    doc.add_paragraph("• 04_download_hospital_compare.py: Pulls facility quality metrics from CMS Hospital Compare.", style='List Bullet')

    # Module 2 & 3: ETL & Database
    doc.add_heading('Module 2 & 3: Transformation & Database Loading', level=2)
    doc.add_paragraph("• 05_clean_and_validate.py: Cleans raw CSVs, standardizes naming conventions, formats ZIP codes, and handles missing values.", style='List Bullet')
    doc.add_paragraph("• 06_load_to_sqlite.py: Acts as the Data Warehouse. Builds a relational SQLite database (urology_analytics.db) for secure, scalable SQL querying.", style='List Bullet')

    # Module 4: Core Analytics
    doc.add_heading('Module 4: Core Analytics Engine', level=2)
    doc.add_paragraph("• 07_analysis_productivity.py: Calculates Relative Value Unit (RVU) generation per physician and calculates revenue estimates.", style='List Bullet')
    doc.add_paragraph("• 08_analysis_billing.py: Analyzes CPT coding frequencies to flag statistical outliers (Billing Anomalies) that represent compliance audit risks.", style='List Bullet')
    doc.add_paragraph("• 09_analysis_capacity.py: Compares clinical volume against provider supply to identify structural capacity bottlenecks.", style='List Bullet')
    doc.add_paragraph("• 10_analysis_benchmarking.py: Compares local department performance against national median and 90th percentile top-performers.", style='List Bullet')

    # Module 5 & 6: Predictive & Scenario Modeling
    doc.add_heading('Module 5: Advanced Modeling', level=2)
    doc.add_paragraph("• 11_predictive_alerts.py: Acts as an operational 'Command Center'. Sweeps all analytics outputs and generates RED/YELLOW flag alerts for productivity drops or capacity grids.", style='List Bullet')
    doc.add_paragraph("• 12_forecasting_arima.py: Utilizes an ARIMAX statistical time-series model to forecast future Medicare revenue over a 6-month horizon.", style='List Bullet')
    doc.add_paragraph("• 13_scenario_modeling.py: Evaluates 'What-If' scenarios (e.g., adding 2 new surgeons, 5% Medicare cuts) to show dynamic financial impact.", style='List Bullet')

    # Module 7: Exporting & Reporting
    doc.add_heading('Module 6: BI & Automated Reporting', level=2)
    doc.add_paragraph("• 14_generate_powerbi_exports.py: Stages pristine CSV cuts and documentation for automated Power BI dashboard consumption.", style='List Bullet')
    doc.add_paragraph("• 15_generate_executive_memo.py: Auto-generates a strategic, formatted Word Document summarizing top performance KPIs and recommendations.", style='List Bullet')
    doc.add_paragraph("• 16_audit_log.py: Compiles an immutable system audit trail to ensure data governance and processing compliance.", style='List Bullet')

    # --- Section 3: LLM Integration ---
    doc.add_heading('3. Artificial Intelligence (LLM) Integration', level=1)
    
    doc.add_heading('Which LLM We Are Using', level=2)
    doc.add_paragraph(
        "We are utilizing OpenAI's gpt-3.5-turbo via the official openai Python package. "
        "This model was chosen for its excellent balance of speed, low latency, and strong reasoning capabilities "
        "when provided with structured data contexts."
    )

    doc.add_heading('How the LLM Works Within the System', level=2)
    p = doc.add_paragraph(
        "The LLM acts as an 'Intelligent Diagnostics Agent' within the predictive alert engine. "
        "Instead of relying on hardcoded, static text to tell executives what to do when a system alert triggers, "
        "we connected the LLM directly into 11_predictive_alerts.py."
    )
    
    doc.add_paragraph("The mechanism operates as follows:", style='List Number')
    doc.add_paragraph("Environment Setup: The system securely loads an OPENAI_API_KEY from a local .env file using the python-dotenv library.", style='List Bullet')
    doc.add_paragraph("Anomaly Detection: The script identifies a business bottleneck (e.g., 'Provider RVUs drop 15% below baseline' or 'Capacity constraints predict 60-day gridlock').", style='List Bullet')
    doc.add_paragraph("Dynamic Prompting: The script constructs a custom prompt injecting the specific variables: Alert Type, Severity, Entity Involved, Trigger Condition, and Expected Impact.", style='List Bullet')
    doc.add_paragraph("API Execution: The OpenAI API is called. The System Prompt instructs the model to act as a 'Specialized Healthcare Data Analyst and Operations Expert.'", style='List Bullet')
    doc.add_paragraph("Output Generation: The LLM returns a concise, diagnostic hypothesis and a prescriptive operational recommendation (e.g., 'Re-allocate 2 OR blocks from general surgery to oncology to offset wait times.').", style='List Bullet')
    doc.add_paragraph("Dashboard Rendering: The AI response is appended to the alerts_report.csv and dynamically loaded into the custom HTML/JS local web dashboard under a designated '🤖 AI Insight' badge.", style='List Bullet')

    doc.add_paragraph(
        "\nFurthermore, the system includes a graceful degradation protocol: if the API key is missing or network connectivity fails, "
        "the script automatically reverts to standard static recommendations so the analytics pipeline never crashes."
    )

    # Save Document
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Urology_Analytics_Project_Documentation.docx')
    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    create_documentation()
